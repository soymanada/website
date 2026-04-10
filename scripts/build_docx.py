"""
SoyManada — Generador de plantilla Word
Ejecutar con: python3 scripts/build_docx.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, lxml.etree as etree

# ── Colores de marca ──────────────────────────────────────────────
IRIS_500   = RGBColor(0x7B, 0x4D, 0xC8)
IRIS_100   = RGBColor(0xE8, 0xDC, 0xF7)
IVORY      = RGBColor(0xFA, 0xF8, 0xF4)
TEXT_700   = RGBColor(0x2D, 0x27, 0x3D)
TEXT_400   = RGBColor(0x6B, 0x63, 0x80)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
WARN_BG    = RGBColor(0xFF, 0xF3, 0xCD)
WARN_BORDER= RGBColor(0xF0, 0xAD, 0x4E)
TIP_BG     = RGBColor(0xE8, 0xF5, 0xE9)

# ── Helpers XML ──────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Fondo de celda por color hex string (sin #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '12')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), color)
            tcBorders.append(b)
    tcPr.append(tcBorders)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb

def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_watermark(doc, text='soymanada.com'):
    """Marca de agua diagonal sutil en el header usando XML raw."""
    VML_NS   = 'urn:schemas-microsoft-com:vml'
    OFFICE_NS = 'urn:schemas-microsoft-com:office:office'
    WORD_NS  = 'urn:schemas-microsoft-com:office:word'

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run()

        # Construir el XML del shape con lxml directo (evita OxmlElement para namespaces VML)
        pict_xml = f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                              xmlns:v="{VML_NS}"
                              xmlns:o="{OFFICE_NS}"
                              xmlns:wvml="{WORD_NS}">
  <v:shape id="WaterMarkShape" type="#_x0000_t136"
    style="position:absolute;margin-left:0;margin-top:0;
           width:500pt;height:200pt;rotation:315;
           z-index:-251654144;
           mso-position-horizontal:center;
           mso-position-horizontal-relative:margin;
           mso-position-vertical:center;
           mso-position-vertical-relative:margin"
    fillcolor="#7B4DC8" stroked="f" o:allowincell="f">
    <v:fill opacity="0.07"/>
    <v:textpath style="font-family:&apos;Calibri&apos;;font-size:54pt;font-weight:bold" string="{text}"/>
  </v:shape>
</w:pict>'''

        pict_el = etree.fromstring(pict_xml)
        run._r.append(pict_el)

# ── Estilos personalizados ────────────────────────────────────────

def define_styles(doc):
    styles = doc.styles

    # Título principal del recurso
    if 'SM-Title' not in [s.name for s in styles]:
        s = styles.add_style('SM-Title', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = 'Calibri'
        s.font.size = Pt(26)
        s.font.bold = True
        s.font.color.rgb = IRIS_500
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(6)

    # Subtítulo
    if 'SM-Subtitle' not in [s.name for s in styles]:
        s = styles.add_style('SM-Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = 'Calibri'
        s.font.size = Pt(13)
        s.font.color.rgb = TEXT_400
        s.font.italic = True
        s.paragraph_format.space_after = Pt(16)

    # Encabezado de sección
    if 'SM-Heading' not in [s.name for s in styles]:
        s = styles.add_style('SM-Heading', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = 'Calibri'
        s.font.size = Pt(13)
        s.font.bold = True
        s.font.color.rgb = IRIS_500
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after = Pt(4)

    # Texto normal
    if 'SM-Body' not in [s.name for s in styles]:
        s = styles.add_style('SM-Body', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = 'Calibri'
        s.font.size = Pt(11)
        s.font.color.rgb = TEXT_700
        s.paragraph_format.space_after = Pt(6)

    # Bullet
    if 'SM-Bullet' not in [s.name for s in styles]:
        s = styles.add_style('SM-Bullet', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = 'Calibri'
        s.font.size = Pt(11)
        s.font.color.rgb = TEXT_700
        s.paragraph_format.left_indent = Cm(0.8)
        s.paragraph_format.space_after = Pt(3)

    return styles

# ── Bloques de contenido ──────────────────────────────────────────

def add_colored_block(doc, heading, body_lines, bg_hex, border_hex=None, emoji=''):
    """Bloque con fondo coloreado usando tabla de 1 celda."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    if border_hex:
        set_cell_border(cell, top=border_hex, bottom=border_hex, left=border_hex, right=border_hex)
    else:
        remove_cell_borders(cell)

    # Heading de bloque
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(f'{emoji}  {heading}' if emoji else heading)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = TEXT_700
    set_para_spacing(p, before=80, after=40)

    # Líneas del cuerpo
    for line in body_lines:
        bp = cell.add_paragraph()
        r = bp.add_run(line)
        r.font.size = Pt(10.5)
        r.font.name = 'Calibri'
        r.font.color.rgb = TEXT_700
        set_para_spacing(bp, before=0, after=30)

    # Espacio post-tabla
    doc.add_paragraph()

def add_step(doc, number, title, description):
    """Paso numerado con línea decorativa."""
    p = doc.add_paragraph(style='SM-Heading')
    run_num = p.add_run(f'Paso {number}  ')
    run_num.font.color.rgb = IRIS_500
    run_num.font.bold = True
    run_num.font.size = Pt(12)
    run_title = p.add_run(title)
    run_title.font.color.rgb = TEXT_700
    run_title.font.bold = True
    run_title.font.size = Pt(12)

    p2 = doc.add_paragraph(style='SM-Body')
    p2.add_run(description)

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E8DCF7')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_spacing(p, before=80, after=80)

def add_header_band(doc, label):
    """Banda morada arriba de sección."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, '7B4DC8')
    remove_cell_borders(cell)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(label.upper())
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = WHITE
    run.font.letter_spacing = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=60, after=60)
    doc.add_paragraph()

# ── Portada ───────────────────────────────────────────────────────

def build_cover(doc):
    # Espacio superior
    for _ in range(3):
        doc.add_paragraph()

    # Logo / marca
    p_logo = doc.add_paragraph()
    r = p_logo.add_run('🐾  SoyManada')
    r.font.name = 'Calibri'
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = IRIS_500
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_sub_logo = doc.add_paragraph()
    r2 = p_sub_logo.add_run('Recursos para migrantes en Canadá')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    r2.font.color.rgb = TEXT_400
    p_sub_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(2):
        doc.add_paragraph()

    # Línea separadora iris
    add_divider(doc)

    # Título recurso (placeholder)
    p_title = doc.add_paragraph()
    r_t = p_title.add_run('[TÍTULO DEL RECURSO]')
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(28)
    r_t.font.bold = True
    r_t.font.color.rgb = IRIS_500
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_subtitle = doc.add_paragraph()
    r_s = p_subtitle.add_run('[Subtítulo breve — una línea que describe el contenido]')
    r_s.font.name = 'Calibri'
    r_s.font.size = Pt(13)
    r_s.font.color.rgb = TEXT_400
    r_s.font.italic = True
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(2):
        doc.add_paragraph()

    add_divider(doc)

    # Pie de portada
    p_foot = doc.add_paragraph()
    r_f = p_foot.add_run('soymanada.com  ·  Versión: [FECHA]  ·  [PAÍS/PROVINCIA]')
    r_f.font.name = 'Calibri'
    r_f.font.size = Pt(9)
    r_f.font.color.rgb = TEXT_400
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Salto de página
    doc.add_page_break()

# ── Plantilla base ────────────────────────────────────────────────

def build_template(doc):
    # 1. Encabezado de sección
    add_header_band(doc, '¿Para quién es este recurso?')

    add_colored_block(
        doc,
        'Este recurso te puede servir si…',
        [
            '▸  [Condición 1 — ej: eres migrante recién llegado a Canadá]',
            '▸  [Condición 2 — ej: quieres saber cómo obtener tu licencia]',
            '▸  [Condición 3 — agrega o elimina según el tema]',
        ],
        'EAD9F5',  # iris claro
    )

    # 2. Qué vas a necesitar
    add_header_band(doc, 'Qué vas a necesitar')

    add_colored_block(
        doc,
        'Documentos y requisitos',
        [
            '□  [Documento o requisito 1]',
            '□  [Documento o requisito 2]',
            '□  [Documento o requisito 3 — agrega los que apliquen]',
        ],
        'F5F5F5',
    )

    # 3. Pasos
    add_header_band(doc, 'Paso a paso')

    for i in range(1, 6):
        add_step(doc,
            number=i,
            title=f'[Título del paso {i}]',
            description=f'Descripción del paso {i}. Explica qué hacer, dónde ir, qué decir o qué enviar. '
                         f'Usa lenguaje simple y directo. Borra este texto y escribe el real.'
        )
        if i < 5:
            add_divider(doc)

    doc.add_paragraph()

    # 4. Tips
    add_colored_block(
        doc,
        'Tips y consejos prácticos',
        [
            '💡  [Tip 1 — algo que facilita el proceso]',
            '💡  [Tip 2 — ahorra tiempo o dinero]',
            '💡  [Tip 3 — error frecuente que se puede evitar]',
        ],
        'E8F5E9',  # verde suave
        emoji='',
    )

    # 5. Advertencia
    add_colored_block(
        doc,
        'Importante / Depende de tu provincia',
        [
            '⚠️   Este contenido es orientativo. La información puede variar según la provincia o cambiar con el tiempo.',
            '⚠️   Siempre verifica en el sitio oficial del gobierno o consulta con un proveedor verificado.',
            '⚠️   [Agrega cualquier advertencia específica del tema aquí]',
        ],
        'FFF3CD',  # amarillo suave
        border_hex='F0AD4E',
    )

    # 6. CTA final
    add_header_band(doc, 'Próximos pasos')

    add_colored_block(
        doc,
        '¿Qué hacer después de leer esto?',
        [
            '→  Explora proveedores verificados en soymanada.com',
            '→  Sigue revisando otros recursos en la sección Primeros Pasos',
            '→  Únete a la comunidad en el grupo de WhatsApp de SoyManada',
            '→  ¿Tienes dudas? Escríbele directamente a un proveedor',
        ],
        'EAD9F5',
    )

# ── Footer ────────────────────────────────────────────────────────

def set_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r1 = p.add_run('🐾 SoyManada  ·  soymanada.com')
        r1.font.name = 'Calibri'
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = IRIS_500
        r1.font.bold = True

        r2 = p.add_run('  ·  Última revisión: [FECHA]  ·  Este documento es orientativo, no constituye asesoría legal.')
        r2.font.name = 'Calibri'
        r2.font.size = Pt(8)
        r2.font.color.rgb = TEXT_400

# ── Ejemplo aplicado: Licencia de conducir ───────────────────────

def build_example(doc):
    doc.add_page_break()

    # Portada del ejemplo
    p = doc.add_paragraph()
    r = p.add_run('EJEMPLO APLICADO')
    r.font.name = 'Calibri'
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = IRIS_500
    r.font.letter_spacing = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    r2 = p2.add_run('Cómo obtener tu licencia de conducir en Canadá')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(22)
    r2.font.bold = True
    r2.font.color.rgb = IRIS_500
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p3 = doc.add_paragraph()
    r3 = p3.add_run('Guía práctica para migrantes con licencia extranjera')
    r3.font.name = 'Calibri'
    r3.font.size = Pt(12)
    r3.font.color.rgb = TEXT_400
    r3.font.italic = True
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_divider(doc)

    # Sirve si...
    add_colored_block(
        doc,
        'Este recurso te puede servir si…',
        [
            '▸  Llegaste a Canadá con licencia de conducir de otro país',
            '▸  Quieres saber si puedes canjearla o si necesitas hacer exámenes',
            '▸  No sabes por dónde empezar en tu provincia',
        ],
        'EAD9F5',
    )

    # Documentos
    add_colored_block(
        doc,
        'Qué vas a necesitar (ejemplo Ontario)',
        [
            '□  Licencia de conducir original del país de origen (vigente o no vencida al llegar)',
            '□  Traducción oficial al inglés (si no está en inglés o francés)',
            '□  Pasaporte vigente',
            '□  Prueba de residencia en Ontario (factura de servicios, contrato)',
            '□  Pago de la tarifa de licencia (varía según categoría)',
        ],
        'F5F5F5',
    )

    # Pasos
    add_header_band(doc, 'Paso a paso — Ontario (G license)')

    steps = [
        ('Verifica si tu licencia es canjeable',
         'Canadá tiene acuerdos de reciprocidad con varios países. Si vienes de México, Colombia, Chile, Argentina u otros países de AL, '
         'es probable que debas hacer exámenes igual. Consulta la lista oficial del MTO (Ministry of Transportation Ontario) antes de ir.'),
        ('Traduce tu licencia si es necesario',
         'Si tu licencia no está en inglés o francés, necesitas una traducción certificada por un traductor acreditado (ATIO en Ontario). '
         'No vale cualquier traducción. [VERIFICAR: requisitos pueden variar — ver ontario.ca]'),
        ('Ve a un DriveTest Centre',
         'Lleva todos tus documentos al centro DriveTest más cercano. Puedes buscar ubicaciones en drivetest.ca. '
         'Pide turno con anticipación, los tiempos de espera pueden ser de varias semanas.'),
        ('Rinde el examen teórico (G1)',
         'Si tu país no tiene convenio de reciprocidad, empezarás desde el nivel G1: un examen escrito de reglas de tránsito. '
         'Puedes prepararte con el manual oficial "The Official MTO Driver\'s Handbook" (gratis en línea).'),
        ('Avanza en el sistema gradual (G1 → G2 → G)',
         'Ontario usa un sistema de licencia gradual (GDL). G1: solo teórico. G2: examen de manejo en vía. G: examen de manejo en autopista. '
         'Los tiempos mínimos entre etapas varían. [VERIFICAR fechas vigentes en ontario.ca]'),
    ]

    for i, (title, desc) in enumerate(steps, 1):
        add_step(doc, i, title, desc)
        if i < len(steps):
            add_divider(doc)

    doc.add_paragraph()

    add_colored_block(
        doc,
        'Tips prácticos',
        [
            '💡  Reserva el turno en DriveTest con semanas de anticipación — los cupos se agotan rápido.',
            '💡  Lleva copias de todos tus documentos además de los originales.',
            '💡  Si vienes de una provincia con acuerdo (ej: Québec ↔ Ontario), el proceso puede ser más directo.',
            '💡  La licencia de otro país puede ser válida los primeros 60 días mientras tramitas la canadiense.',
        ],
        'E8F5E9',
    )

    add_colored_block(
        doc,
        'Importante — verifica antes de actuar',
        [
            '⚠️   Las reglas cambian por provincia. Este ejemplo es orientativo para Ontario.',
            '⚠️   Si vienes de un país con convenio de reciprocidad, el proceso puede ser diferente.',
            '⚠️   Siempre verifica en ontario.ca/en/driving o el sitio oficial de tu provincia.',
            '⚠️   Esta guía no reemplaza asesoría de un profesional o el sitio oficial del gobierno.',
        ],
        'FFF3CD',
        border_hex='F0AD4E',
    )

    add_colored_block(
        doc,
        '¿Qué hacer después?',
        [
            '→  Busca proveedores de servicios de traducción certificada en soymanada.com',
            '→  Revisa la guía de "Primeros pasos en tu provincia" en SoyManada',
            '→  Únete a la comunidad — otros migrantes ya pasaron por esto',
        ],
        'EAD9F5',
    )

# ── Nota de uso ───────────────────────────────────────────────────

def build_usage_note(doc):
    doc.add_page_break()

    p = doc.add_paragraph()
    r = p.add_run('Cómo usar esta plantilla')
    r.font.name = 'Calibri'
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = IRIS_500

    items = [
        ('Duplicar el archivo',
         'Crea una copia del .docx antes de editar. Nunca trabajes sobre el original.'),
        ('Reemplazar los placeholders',
         'Busca los textos entre corchetes [COMO ESTE] y reemplázalos con el contenido real. '
         'Usa Ctrl+H para reemplazar en lote.'),
        ('Usar los estilos SM-*',
         'El documento tiene estilos propios: SM-Title, SM-Subtitle, SM-Heading, SM-Body, SM-Bullet. '
         'Úsalos desde el panel de Estilos de Word para mantener consistencia.'),
        ('Bloques de color',
         'Los bloques de consejos, advertencias y CTA son tablas de 1 celda. '
         'Para duplicar uno: selecciónalo, cópialo y pégalo donde necesites.'),
        ('Marca de agua',
         'La marca de agua "soymanada.com" está en el encabezado. '
         'Para quitarla: Ve a Insertar → Encabezado → Editar encabezado → borra el elemento.'),
        ('Exportar a PDF',
         'Archivo → Exportar → Crear PDF/XPS. El PDF queda con el mismo diseño.'),
    ]

    for title, body in items:
        p_h = doc.add_paragraph(style='SM-Heading')
        p_h.add_run(title)
        p_b = doc.add_paragraph(style='SM-Body')
        p_b.add_run(body)

# ── Main ──────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    define_styles(doc)
    add_watermark(doc, 'soymanada.com')
    set_footer(doc)

    build_cover(doc)
    build_template(doc)
    build_example(doc)
    build_usage_note(doc)

    out_path = '/home/user/website/public/downloads/soymanada-plantilla-primeros-pasos.docx'
    import os; os.makedirs('/home/user/website/public/downloads', exist_ok=True)
    doc.save(out_path)
    print(f'✅  Documento generado: {out_path}')

if __name__ == '__main__':
    main()
