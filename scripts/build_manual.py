"""
SoyManada — Generador de manuales Word (Migrante + Proveedor)
Ejecutar con: python3 scripts/build_manual.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import lxml.etree as etree

# ── Colores de marca ──────────────────────────────────────────────
IRIS_500  = RGBColor(0x7B, 0x4D, 0xC8)
IVORY     = RGBColor(0xFA, 0xF8, 0xF4)
TEXT_700  = RGBColor(0x2D, 0x27, 0x3D)
TEXT_400  = RGBColor(0x6B, 0x63, 0x80)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

WARN_BG     = 'FFF3CD'
WARN_BORDER = 'F0AD4E'
TIP_BG      = 'E8F5E9'

DATE_STR      = 'Abril 2026'
DATE_STR_V2   = 'Mayo 2026'
IMG_PLACEHOLDER_BG = 'EDE9F8'   # lavanda suave para recuadros de imagen

# ── Helpers XML ──────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '12')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), color)
            tcBorders.append(b)
    tcPr.append(tcBorders)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E8DCF7')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_spacing(p, before=80, after=80)

# ── Watermark ────────────────────────────────────────────────────

def add_watermark(doc, text='soymanada.com'):
    """Marca de agua diagonal sutil en el header usando lxml directo (no OxmlElement para VML)."""
    VML_NS    = 'urn:schemas-microsoft-com:vml'
    OFFICE_NS = 'urn:schemas-microsoft-com:office:office'
    WORD_NS   = 'urn:schemas-microsoft-com:office:word'

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run()

        pict_xml = (
            '<w:pict'
            ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' xmlns:v="{VML_NS}"'
            f' xmlns:o="{OFFICE_NS}"'
            f' xmlns:wvml="{WORD_NS}">'
            '<v:shape'
            ' style="position:absolute;width:500pt;height:200pt;rotation:315;'
            'z-index:-251654144;'
            'mso-position-horizontal:center;'
            'mso-position-horizontal-relative:margin;'
            'mso-position-vertical:center;'
            'mso-position-vertical-relative:margin"'
            ' fillcolor="#7B4DC8" stroked="f">'
            '<v:fill opacity="0.07"/>'
            '<v:textpath'
            " style=\"font-family:'Calibri';font-size:54pt;font-weight:bold\""
            f' string="{text}"/>'
            '</v:shape>'
            '</w:pict>'
        )

        pict_el = etree.fromstring(pict_xml)
        run._r.append(pict_el)

# ── Footer ────────────────────────────────────────────────────────

def set_footer(doc, date_str=DATE_STR):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r1 = p.add_run('🐾 SoyManada · soymanada.com')
        r1.font.name = 'Calibri'
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = IRIS_500
        r1.font.bold = True

        r2 = p.add_run(f'  ·  {date_str}  ·  Documento orientativo, no constituye asesoría.')
        r2.font.name = 'Calibri'
        r2.font.size = Pt(8)
        r2.font.color.rgb = TEXT_400

# ── Estilos ───────────────────────────────────────────────────────

def define_styles(doc):
    styles = doc.styles
    existing = [s.name for s in styles]

    if 'SM-Title' not in existing:
        s = styles.add_style('SM-Title', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = 'Calibri'
        s.font.size = Pt(26)
        s.font.bold = True
        s.font.color.rgb = IRIS_500
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(6)

    if 'SM-Subtitle' not in existing:
        s = styles.add_style('SM-Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = 'Calibri'
        s.font.size = Pt(13)
        s.font.color.rgb = TEXT_400
        s.font.italic = True
        s.paragraph_format.space_after = Pt(16)

    if 'SM-Heading' not in existing:
        s = styles.add_style('SM-Heading', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = 'Calibri'
        s.font.size = Pt(13)
        s.font.bold = True
        s.font.color.rgb = IRIS_500
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after = Pt(4)

    if 'SM-Body' not in existing:
        s = styles.add_style('SM-Body', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = 'Calibri'
        s.font.size = Pt(11)
        s.font.color.rgb = TEXT_700
        s.paragraph_format.space_after = Pt(6)

    if 'SM-Bullet' not in existing:
        s = styles.add_style('SM-Bullet', WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = 'Calibri'
        s.font.size = Pt(11)
        s.font.color.rgb = TEXT_700
        s.paragraph_format.left_indent = Cm(0.8)
        s.paragraph_format.space_after = Pt(3)

# ── Bloques de contenido ──────────────────────────────────────────

def add_colored_block(doc, heading, body_lines, bg_hex, border_hex=None, emoji=''):
    """Bloque con fondo coloreado usando tabla de 1 celda sin bordes."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    if border_hex:
        set_cell_border(cell, top=border_hex, bottom=border_hex, left=border_hex, right=border_hex)
    else:
        remove_cell_borders(cell)

    # Heading del bloque
    p = cell.paragraphs[0]
    p.clear()
    label = f'{emoji}  {heading}' if emoji else heading
    run = p.add_run(label)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = TEXT_700
    set_para_spacing(p, before=80, after=40)

    # Líneas del cuerpo
    for line in body_lines:
        bp = cell.add_paragraph()
        r = bp.add_run(line)
        r.font.size = Pt(10.5)
        r.font.name = 'Calibri'
        r.font.color.rgb = TEXT_700
        set_para_spacing(bp, before=0, after=30)

    doc.add_paragraph()  # spacing after

def add_numbered_steps(doc, steps):
    """Lista numerada: cada step es un string. Número en negrita."""
    for i, step_text in enumerate(steps, 1):
        p = doc.add_paragraph(style='SM-Body')
        p.paragraph_format.left_indent = Cm(0.5)
        run_num = p.add_run(f'{i}.  ')
        run_num.font.bold = True
        run_num.font.name = 'Calibri'
        run_num.font.color.rgb = IRIS_500
        run_body = p.add_run(step_text)
        run_body.font.name = 'Calibri'
        run_body.font.color.rgb = TEXT_700

def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style='SM-Bullet')
        run = p.add_run(f'▸  {item}')
        run.font.name = 'Calibri'
        run.font.color.rgb = TEXT_700

def add_section_heading(doc, title):
    p = doc.add_paragraph(style='SM-Heading')
    p.add_run(title)

def add_body(doc, text):
    p = doc.add_paragraph(style='SM-Body')
    p.add_run(text)

def add_tip_block(doc, tip_text):
    add_colored_block(doc, 'Consejo', [tip_text], TIP_BG)

def add_warning_block(doc, warning_text):
    add_colored_block(doc, 'Importante', [warning_text], WARN_BG, border_hex=WARN_BORDER)

def add_faq_item(doc, question, answer):
    p_q = doc.add_paragraph(style='SM-Body')
    r_q = p_q.add_run(question)
    r_q.font.bold = True
    r_q.font.name = 'Calibri'
    r_q.font.color.rgb = TEXT_700

    p_a = doc.add_paragraph(style='SM-Body')
    p_a.paragraph_format.left_indent = Cm(0.5)
    r_a = p_a.add_run(answer)
    r_a.font.name = 'Calibri'
    r_a.font.color.rgb = TEXT_400
    set_para_spacing(p_a, before=0, after=120)

def add_screenshot_placeholder(doc, label, height_cm=5.5):
    """Recuadro punteado lavanda que marca dónde va un screenshot."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, IMG_PLACEHOLDER_BG)
    set_cell_border(cell, top='7B4DC8', bottom='7B4DC8', left='7B4DC8', right='7B4DC8')
    # altura mínima aproximada
    from docx.oxml.ns import qn as _qn
    tc = cell._tc
    trPr = tc.getparent().get_or_add_trPr() if hasattr(tc.getparent(), 'get_or_add_trPr') else None
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=int(height_cm * 100), after=int(height_cm * 100))
    r = p.add_run(f'📷  {label}')
    r.font.name    = 'Calibri'
    r.font.size    = Pt(10.5)
    r.font.italic  = True
    r.font.color.rgb = IRIS_500
    doc.add_paragraph()

# ── Portada genérica ──────────────────────────────────────────────

def build_cover(doc, title, subtitle, footer_line):
    for _ in range(3):
        doc.add_paragraph()

    p_logo = doc.add_paragraph()
    r = p_logo.add_run('🐾 SoyManada')
    r.font.name = 'Calibri'
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = IRIS_500
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(2):
        doc.add_paragraph()

    add_divider(doc)

    p_title = doc.add_paragraph()
    r_t = p_title.add_run(title)
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(28)
    r_t.font.bold = True
    r_t.font.color.rgb = IRIS_500
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_subtitle = doc.add_paragraph()
    r_s = p_subtitle.add_run(subtitle)
    r_s.font.name = 'Calibri'
    r_s.font.size = Pt(13)
    r_s.font.color.rgb = TEXT_400
    r_s.font.italic = True
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(2):
        doc.add_paragraph()

    add_divider(doc)

    p_foot = doc.add_paragraph()
    r_f = p_foot.add_run(footer_line)
    r_f.font.name = 'Calibri'
    r_f.font.size = Pt(9)
    r_f.font.color.rgb = TEXT_400
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

# ── Tabla de planes (proveedor) ───────────────────────────────────

def add_plans_table(doc):
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    headers = ['Plan', 'Precio', 'Incluye']
    rows_data = [
        ('Bronze (Gratis)', '$0', 'Perfil básico, 1 categoría, mensajería, reseñas'),
        ('Activa (Silver)', '$4.990 CLP/mes', 'Todo Bronze + métricas, calendario de citas, WhatsApp visible, posición prioritaria, responder reseñas'),
        ('Pro (Gold)', '$14.990 CLP/mes', 'Todo Activa + todas las categorías, métricas completas, herramientas avanzadas'),
    ]

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '7B4DC8')
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, (plan, precio, incluye) in enumerate(rows_data, 1):
        row = table.rows[r_idx]
        bg = 'EAD9F5' if r_idx % 2 == 0 else 'F5F0FF'
        for c_idx, val in enumerate([plan, precio, incluye]):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_700

    doc.add_paragraph()

# ── MANUAL MIGRANTE ───────────────────────────────────────────────

def build_manual_migrante(doc):
    # Portada
    build_cover(
        doc,
        title='Manual del usuario — Migrante',
        subtitle='Guía completa para encontrar servicios, contactar proveedores y aprovechar la comunidad.',
        footer_line='soymanada.com · Versión 1.0 · Abril 2026',
    )

    # ── Sección 1: Registro y acceso ──────────────────────────────
    add_section_heading(doc, '1. Crear tu cuenta')
    add_numbered_steps(doc, [
        'Ir a soymanada.com → botón "Iniciar sesión".',
        'Elegir Google (recomendado) o email + contraseña.',
        'Si usas email: recibirás un correo de confirmación — debes hacer click antes de entrar.',
        'Una vez confirmado, la sesión se mantiene activa.',
    ])
    add_tip_block(doc, 'El registro es gratuito y no necesitas tarjeta de crédito.')

    # ── Sección 2: Explorar proveedores ──────────────────────────
    add_section_heading(doc, '2. Encontrar el proveedor correcto')
    add_body(doc,
        'El directorio en /proveedores te permite explorar todos los proveedores disponibles. '
        'Puedes filtrar por categoría, país y activar la opción de "solo verificados". '
        'La búsqueda por texto filtra nombre, servicio y descripción del proveedor.'
    )
    add_tip_block(doc, "Activa 'Solo verificados' para ver solo proveedores que pasaron la revisión manual de SoyManada.")
    add_numbered_steps(doc, [
        'Usa el filtro de categoría.',
        'Escribe un término en el buscador.',
        'Haz click en la tarjeta para ver el perfil completo.',
    ])

    # ── Sección 3: Perfil de proveedor ────────────────────────────
    add_section_heading(doc, '3. Qué ver en el perfil de un proveedor')
    add_bullet_list(doc, [
        'Nombre y servicio',
        'Badge verificado (si aplica)',
        'Descripción del servicio',
        'Países donde opera',
        'Idiomas',
        'Reseñas de otros migrantes',
        'Botón de contacto',
        'WhatsApp (si el proveedor lo activó)',
    ])
    add_warning_block(doc,
        'El badge verificado indica que SoyManada revisó al proveedor, no garantiza el resultado de su servicio. '
        'Lee siempre las reseñas.'
    )

    # ── Sección 4: Contactar a un proveedor ──────────────────────
    add_section_heading(doc, '4. Enviar un mensaje')
    add_numbered_steps(doc, [
        'Debes estar con sesión iniciada.',
        'Clic en "Enviar mensaje" en el sidebar del perfil.',
        'Escribe tu consulta (máximo 1000 caracteres).',
        'Clic en "Enviar mensaje".',
        'Recibirás respuesta por email cuando el proveedor conteste.',
    ])
    add_body(doc,
        'Si el proveedor tiene WhatsApp o Instagram activados, también verás esos botones de contacto en su perfil.'
    )
    add_tip_block(doc,
        'Sé específico en tu mensaje: provincia, tipo de visa, servicio que buscas. '
        'Así el proveedor puede darte una respuesta más útil.'
    )

    # ── Sección 5: Dejar una reseña ──────────────────────────────
    add_section_heading(doc, '5. Evaluar a un proveedor')
    add_numbered_steps(doc, [
        'Entra al perfil del proveedor.',
        'Clic en "Dejar una evaluación" (debajo del sidebar).',
        'Selecciona entre 1 y 5 patas (rating general).',
        'Opcionalmente evalúa subcategorías: velocidad, confiabilidad, claridad, valor.',
        'Escribe un comentario (opcional, máx. 300 caracteres).',
        'Clic en "Enviar evaluación". Solo puedes dejar una reseña por proveedor.',
    ])
    add_tip_block(doc, 'Tu reseña ayuda al próximo migrante. Sé honesto y específico.')

    # ── Sección 6: Primeros pasos ─────────────────────────────────
    add_section_heading(doc, '6. La sección Primeros Pasos')
    add_body(doc,
        'La sección /primeros-pasos contiene guías sobre temas clave para recién llegados: '
        'SIN Number, banca, arriendo, trabajo y visas en Canadá. '
        'Es información orientativa pensada para orientarte — siempre verifica los datos en los sitios oficiales correspondientes.'
    )
    add_warning_block(doc,
        'Esta información es orientativa. Para decisiones migratorias, consulta siempre el sitio oficial '
        'del gobierno canadiense (canada.ca) o un asesor autorizado.'
    )

    # ── Sección 7: Solicitar un servicio ──────────────────────────
    add_section_heading(doc, '7. No encuentras lo que buscas')
    add_body(doc,
        'En la cuadrícula de categorías encontrarás dos botones especiales: '
        '"¿No ves tu categoría?" te permite sugerir una nueva categoría, y '
        '"¿Buscas algo específico?" te permite solicitar un servicio concreto. '
        'El equipo de SoyManada revisa todas las solicitudes y te avisa si encuentran un proveedor adecuado.'
    )

    # ── Sección 8: Preguntas frecuentes ──────────────────────────
    add_section_heading(doc, '8. Preguntas frecuentes')
    doc.add_paragraph()

    faqs = [
        ('¿Es gratis usar SoyManada?',
         'Sí, el acceso al directorio y los mensajes son completamente gratuitos para migrantes.'),
        ('¿Los proveedores están verificados?',
         'SoyManada revisa manualmente a cada proveedor antes de publicarlo. El badge verificado confirma que pasó ese proceso.'),
        ('¿Mis datos están seguros?',
         'SoyManada usa Supabase (infraestructura de nivel enterprise) para almacenar tus datos. '
         'Tu email nunca se comparte con proveedores sin tu consentimiento.'),
        ('¿Puedo editar o borrar mi reseña?',
         'Por ahora solo puedes editar una reseña existente. Contacta al equipo si necesitas borrarla.'),
        ('¿Hay app móvil?',
         'Aún no. El sitio web está optimizado para móvil y funciona bien desde el navegador de tu celular.'),
    ]
    for q, a in faqs:
        add_faq_item(doc, q, a)

# ── MANUAL PROVEEDOR ──────────────────────────────────────────────

def build_manual_proveedor(doc):
    # ── Portada ───────────────────────────────────────────────────
    build_cover(
        doc,
        title='Manual del proveedor',
        subtitle='Guía completa para gestionar tu perfil, mensajes, reseñas, pagos y visibilidad en el directorio.',
        footer_line='soymanada.com · Versión 2.0 · Mayo 2026',
    )

    # ── Sección 1: Alta como proveedor ───────────────────────────
    add_section_heading(doc, '1. Cómo darte de alta como proveedor')
    add_body(doc,
        'El proceso de registro es simple y está completamente en línea. '
        'Después de enviar tu solicitud recibirás dos emails: uno inmediato de confirmación y, '
        'una vez aprobado por el equipo, un email de bienvenida con tus instrucciones de acceso.'
    )
    add_numbered_steps(doc, [
        'Ve a soymanada.com/registro-proveedores.',
        'Completa el formulario: nombre del negocio, servicio, categorías, descripción, '
        'idiomas atendidos, países, WhatsApp, email de contacto.',
        'Acepta los términos y envía la solicitud.',
        'Recibirás un email de confirmación inmediatamente — revisa tu bandeja de spam.',
        'El email incluye el link para crear tu cuenta en SoyManada con ese mismo email.',
        'El equipo revisará tu solicitud (1–2 días hábiles).',
        'Si es aprobado, recibirás un email de bienvenida con el enlace a tu dashboard.',
    ])
    add_tip_block(doc,
        'Una descripción detallada y un WhatsApp activo aumentan tu probabilidad de aprobación y tu '
        'visibilidad en búsquedas. Usa el mismo email en el formulario y en tu cuenta SoyManada — '
        'de lo contrario, el sistema no podrá vincular automáticamente tu perfil.'
    )
    add_screenshot_placeholder(doc, 'SCREENSHOT: Página /registro-proveedores — formulario completo visible', height_cm=6)

    add_divider(doc)

    # ── Sección 2: Early Bird — Trial Gold 3 meses gratis ─────────
    add_section_heading(doc, '2. 🚀 Early Bird: 3 meses de Gold gratis')
    add_body(doc,
        'Como parte del lanzamiento de la plataforma, todos los proveedores que se registren '
        'antes del 30 de junio de 2026 tienen acceso a un trial de 3 meses en plan Gold (Pro) sin costo. '
        'Esto incluye todas las funciones avanzadas: métricas completas, calendario de citas, '
        'WhatsApp visible, posición prioritaria y acceso a todas las categorías.'
    )
    add_colored_block(doc,
        '¿Cómo activar el trial?',
        [
            '1. Inicia sesión en tu cuenta de proveedor.',
            '2. Ve a /mi-perfil → tab "Mi Plan".',
            '3. Verás el banner Early Bird con el botón "Activar 3 meses Gold gratis".',
            '4. Clic en el botón — se activa de inmediato, sin tarjeta de crédito.',
            '5. Tu plan pasará a Gold durante los próximos 3 meses.',
        ],
        'EDE9F8',
    )
    add_screenshot_placeholder(doc, 'SCREENSHOT: Banner Early Bird en el tab "Mi Plan" — botón de activación visible', height_cm=4)
    add_warning_block(doc,
        'El trial Early Bird está disponible solo para registros hasta el 30 de junio de 2026. '
        'Al terminar el trial, el plan vuelve a Bronze automáticamente. '
        'Si registros después de esa fecha, el trial será de 1 mes.'
    )

    add_divider(doc)

    # ── Sección 3: Editar tu perfil ───────────────────────────────
    add_section_heading(doc, '3. Dashboard — Tab Perfil')
    add_body(doc, 'Desde el tab "Perfil" puedes editar toda la información que verán los migrantes:')
    add_bullet_list(doc, [
        'Nombre del negocio y título del servicio',
        'Descripción (aparece completa en tu perfil público)',
        'Idiomas en que atiendes',
        'Número de WhatsApp de contacto',
        'Link de Instagram y sitio web',
        'Link de pago (si ofreces pagos online)',
        'Foto de perfil (JPG/PNG, máx 2 MB)',
    ])
    add_numbered_steps(doc, [
        'Inicia sesión → /mi-perfil.',
        'Edita los campos deseados en el tab "Perfil".',
        'Clic en "Guardar cambios".',
    ])
    add_screenshot_placeholder(doc, 'SCREENSHOT: Tab Perfil del dashboard — formulario de edición', height_cm=5.5)

    add_section_heading(doc, 'Badge "Verificado por Manada"')
    add_body(doc,
        'El badge de verificación indica que SoyManada revisó manualmente a este proveedor y validó '
        'su identidad y servicio. Cuando tu perfil es verificado, el badge aparece en tu tarjeta en '
        'el directorio y en tu perfil público, generando mayor confianza en los migrantes.'
    )
    add_tip_block(doc,
        'La verificación es realizada por el equipo de SoyManada. Si deseas solicitarla, escribe a '
        'hola@soymanada.com con tu nombre de negocio y descripción de servicio.'
    )
    add_screenshot_placeholder(doc, 'SCREENSHOT: Tarjeta de proveedor en el directorio con badge "Verificado" visible', height_cm=3.5)

    add_divider(doc)

    # ── Sección 4: Mensajes ───────────────────────────────────────
    add_section_heading(doc, '4. Tab Mensajes — Bandeja de entrada')
    add_body(doc,
        'El tab Mensajes muestra todas las conversaciones con migrantes. '
        'Selecciona un hilo para ver el historial completo y responder directamente desde el dashboard. '
        'Usa Ctrl+Enter para enviar tu respuesta rápidamente.'
    )
    add_numbered_steps(doc, [
        'Tab "Mensajes" en el dashboard.',
        'Selecciona una conversación de la lista.',
        'Lee el mensaje completo y escribe tu respuesta.',
        'Ctrl+Enter o clic en "Responder".',
    ])
    add_screenshot_placeholder(doc, 'SCREENSHOT: Tab Mensajes — lista de conversaciones + hilo abierto', height_cm=5.5)

    add_section_heading(doc, 'Notificaciones por email')
    add_body(doc,
        'En el tab "Perfil" encontrarás la sección de notificaciones con dos toggles: '
        '"Nuevo mensaje" te avisa por email cuando un migrante te escribe. '
        '"Nueva reseña" te avisa cuando alguien deja una evaluación en tu perfil. '
        'Guarda tus preferencias con el botón "Guardar preferencias".'
    )

    add_divider(doc)

    # ── Sección 5: Reseñas verificadas ────────────────────────────
    add_section_heading(doc, '5. Evaluaciones y reseñas verificadas')
    add_body(doc,
        'Las reseñas aparecen en tu perfil público y contribuyen directamente a tu reputación en el directorio. '
        'SoyManada distingue dos tipos de reseña:'
    )
    add_bullet_list(doc, [
        'Reseña estándar — cualquier usuario registrado puede dejar una reseña.',
        'Reseña verificada ✓ — el migrante tuvo una interacción real contigo (mensaje o cita confirmada). '
        'Estas reseñas llevan un badge de verificación y tienen mayor peso en tu puntuación.',
    ])
    add_tip_block(doc,
        'Responder reseñas —incluso las negativas— demuestra profesionalismo. '
        'Los proveedores con plan Activa o Pro (Silver/Gold) pueden publicar respuestas visibles '
        'directamente desde el dashboard. En plan Bronze las reseñas son solo de lectura.'
    )
    add_screenshot_placeholder(doc, 'SCREENSHOT: Perfil público con reseñas — badge "Verificada" en una reseña', height_cm=4)

    add_divider(doc)

    # ── Sección 6: Visibilidad de WhatsApp ────────────────────────
    add_section_heading(doc, '6. Controlar la visibilidad de WhatsApp (Silver/Gold)')
    add_body(doc,
        'En el tab "Herramientas" encontrarás el toggle de visibilidad de WhatsApp. '
        'Cuando está activo, tu número aparece como botón directo en tu perfil público '
        'y los migrantes pueden contactarte con un solo clic.'
    )
    add_warning_block(doc,
        'Esta opción solo está disponible en planes Activa (Silver) y Pro (Gold). '
        'En plan Bronze, el número de WhatsApp NO aparece aunque lo hayas registrado. '
        'Si estás en el trial Gold, tienes acceso completo durante los 3 meses.'
    )

    add_divider(doc)

    # ── Sección 7: Fotos de comunidad ─────────────────────────────
    add_section_heading(doc, '7. Subir fotos a la Comunidad SoyManada')
    add_body(doc,
        'La sección "Historias de la Manada" en la página principal muestra fotos reales de '
        'migrantes en Canadá. Como proveedor o migrante registrado puedes contribuir fotos '
        'que representen tu experiencia de vida en el país.'
    )
    add_numbered_steps(doc, [
        'En la página principal, clic en el botón "Comparte tu foto" bajo el título de la sección.',
        'Selecciona una imagen de tu dispositivo (JPG/PNG, máx 5 MB).',
        'Añade tu nombre y ciudad.',
        'Escribe un mensaje corto opcional.',
        'Clic en "Enviar" — el equipo de SoyManada revisará y publicará la foto.',
    ])
    add_colored_block(doc,
        '¿Qué tipo de fotos funcionan mejor?',
        [
            '▸  Tú en un paisaje o ciudad de Canadá (parques, montañas, barrios).',
            '▸  Tu lugar de trabajo o negocio.',
            '▸  Momentos de integración: comunidad, eventos, vida cotidiana.',
            '▸  Fotos nítidas, bien iluminadas, en horizontal si es posible.',
        ],
        TIP_BG,
    )
    add_screenshot_placeholder(doc, 'SCREENSHOT: Modal "Comparte tu foto" — formulario de subida de imagen', height_cm=4.5)

    add_divider(doc)

    # ── Sección 8: Calendario de citas ───────────────────────────
    add_section_heading(doc, '8. Agenda de citas (Silver/Gold)')
    add_body(doc,
        'El tab "Reservas" permite gestionar citas con migrantes directamente desde el dashboard. '
        'Primero configuras tu disponibilidad semanal; luego, los migrantes pueden agendar una cita '
        'en los horarios libres que hayas definido. Las citas aparecen para que las confirmes o rechaces.'
    )
    add_numbered_steps(doc, [
        'Tab "Reservas" → sección "Gestionar disponibilidad".',
        'Selecciona los días y define el rango horario de cada uno.',
        'Elige la duración de cada cita (ej. 30 o 60 minutos).',
        'Guarda la disponibilidad.',
        'Las citas que lleguen aparecerán en la lista con estado "Pendiente".',
        'Confirma o rechaza cada cita con los botones de acción.',
    ])
    add_warning_block(doc,
        'El calendario de citas solo es visible en tu perfil público si tienes plan Activa o Pro activo. '
        'En Bronze, los migrantes no verán la opción de agendar cita.'
    )

    add_divider(doc)

    # ── Sección 9: Métricas ───────────────────────────────────────
    add_section_heading(doc, '9. Tab Métricas (Silver/Gold)')
    add_body(doc,
        'Las métricas muestran el rendimiento de tu perfil con 6 indicadores clave:'
    )
    add_bullet_list(doc, [
        'Visitas esta semana — cuántas veces se vio tu perfil.',
        'Clics en contacto — cuántas veces alguien hizo clic en WhatsApp, Instagram o mensaje.',
        'Tasa de contacto — % de visitantes que intentaron contactarte.',
        'Puntaje promedio — promedio de tus reseñas (1–5 estrellas).',
        'Mensajes recibidos — total de mensajes en el período.',
        'Tasa de respuesta — % de conversaciones donde respondiste al menos una vez.',
    ])
    add_tip_block(doc,
        'Una tasa de respuesta alta (>80 %) mejora tu posición en el directorio y genera más confianza. '
        'Responde los mensajes aunque sea para decir que no puedes atender — cuenta igual.'
    )

    add_divider(doc)

    # ── Sección 10: Plan y pagos ──────────────────────────────────
    add_section_heading(doc, '10. Tab Mi Plan — Upgrades y pagos')
    add_body(doc, 'Comparativa de planes disponibles:')
    add_plans_table(doc)

    add_body(doc,
        'Para subir de plan, ve al tab "Mi Plan" y haz clic en el botón de upgrade del plan deseado. '
        'El pago se procesa a través de Mercado Pago — puedes pagar con tarjeta de débito, '
        'crédito o transferencia. La suscripción se renueva mensualmente y puedes cancelar en cualquier momento '
        'escribiendo a hola@soymanada.com.'
    )
    add_screenshot_placeholder(doc, 'SCREENSHOT: Tab "Mi Plan" — tabla de planes + botón "Activar Silver" con Mercado Pago', height_cm=5)
    add_tip_block(doc,
        'Si estás en el trial Early Bird, no necesitas pagar nada durante 3 meses. '
        'Puedes contratar un plan pago al terminar el trial o en cualquier momento desde el tab Mi Plan.'
    )

    add_divider(doc)

    # ── Sección 11: Preguntas frecuentes ─────────────────────────
    add_section_heading(doc, '11. Preguntas frecuentes')
    doc.add_paragraph()

    faqs = [
        ('¿Cuándo aparece mi perfil en el directorio?',
         'Después de que el equipo apruebe tu solicitud. El proceso suele tardar 1–2 días hábiles. '
         'Recibirás un email de bienvenida cuando esté listo.'),
        ('¿Por qué no veo mi dashboard después de que me aprobaron?',
         'Verifica que tu cuenta de SoyManada esté registrada con el mismo email que usaste en el formulario. '
         'Si no la creaste aún, usa el link del email de confirmación que recibiste al aplicar.'),
        ('¿Puedo tener más de una categoría?',
         'Con Bronze: 1 categoría. Con Activa (Silver): hasta 2. Con Pro (Gold) o trial: todas las categorías disponibles.'),
        ('¿Los mensajes son privados?',
         'Sí. Solo tú y el migrante que te escribió pueden leer esa conversación.'),
        ('¿El trial Gold Early Bird requiere tarjeta de crédito?',
         'No. Se activa con un clic desde el tab Mi Plan, sin datos de pago. Al terminar el trial, '
         'el plan vuelve a Bronze automáticamente.'),
        ('¿Cómo funciona el pago con Mercado Pago?',
         'Al hacer clic en el botón de upgrade serás redirigido a Mercado Pago. '
         'Puedes pagar con tarjeta de débito, crédito o transferencia. '
         'La suscripción se renueva automáticamente cada mes hasta que la canceles.'),
        ('¿Qué es una reseña verificada?',
         'Una reseña verificada indica que el migrante tuvo una interacción real contigo '
         '(te envió un mensaje o agendó una cita confirmada). Llevan un badge especial en tu perfil.'),
        ('¿Cómo cambio mi contraseña?',
         'Desde la página de login, clic en "¿Olvidaste tu contraseña?" e ingresa tu email.'),
        ('¿Qué pasa si quiero dar de baja mi perfil?',
         'Escribe a hola@soymanada.com con tu nombre de negocio. El equipo lo desactiva en 24–48 horas.'),
    ]
    for q, a in faqs:
        add_faq_item(doc, q, a)

# ── Main ──────────────────────────────────────────────────────────

def make_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    define_styles(doc)
    add_watermark(doc)
    set_footer(doc)
    return doc

def main():
    out_dir = '/home/user/website/public/downloads'
    os.makedirs(out_dir, exist_ok=True)

    # ── Manual Migrante ──────────────────────────────────────────
    doc_m = make_doc()
    build_manual_migrante(doc_m)
    path_m = os.path.join(out_dir, 'manual-migrante-soymanada.docx')
    doc_m.save(path_m)
    size_m = os.path.getsize(path_m)
    print(f'Migrante  → {path_m}  ({size_m:,} bytes)')

    # ── Manual Proveedor ─────────────────────────────────────────
    doc_p = make_doc()
    set_footer(doc_p, date_str=DATE_STR_V2)
    build_manual_proveedor(doc_p)
    path_p = os.path.join(out_dir, 'manual-proveedor-soymanada-v2.docx')
    doc_p.save(path_p)
    size_p = os.path.getsize(path_p)
    print(f'Proveedor → {path_p}  ({size_p:,} bytes)')

    print('Done.')

if __name__ == '__main__':
    main()
